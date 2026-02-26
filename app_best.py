import os
import re
import uuid
import zipfile
import shutil
import subprocess
import calendar
from datetime import datetime, date
from typing import Optional, List, Tuple, Dict

from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename

from docx import Document

# Try to use openpyxl for XLSX output; fall back to CSV if not available
try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter
except Exception:
    openpyxl = None


APP_DIR = os.path.dirname(os.path.abspath(__file__))

UPLOAD_FOLDER = os.path.join(APP_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(APP_DIR, "output")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {"doc", "docx"}

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER

# NOTE: Your original was 30MB. Evidence ZIPs can easily exceed that.
# If you want to keep it at 30MB, change this back.
app.config["MAX_CONTENT_LENGTH"] = 250 * 1024 * 1024  # 250 MB


# -----------------------------
# Regex / Patterns
# -----------------------------
MONTHS = (
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December"
)

DATE_LINE_RE = re.compile(
    r"^DATE:\s+(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(\d{4})\s*$"
)

# Matches "Due Date:" or "DUE DATE:" etc.
DUE_DATE_LINE_RE = re.compile(
    r"^(due date):\s+(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(\d{4})\s*$",
    re.IGNORECASE
)

ENGAGEMENT_YEAR_RE = re.compile(r"\b(20\d{2})(?=\s+SOC Audit engagement\b)")

# Whole-word replacements
FINAL_WORD_RE = re.compile(r"\bfinal\b", re.IGNORECASE)
INITIAL_WORD_RE = re.compile(r"\binitial\b", re.IGNORECASE)

# INVOICE NUMBER: 6 digits
INVOICE_NUMBER_RE = re.compile(r"(INVOICE NUMBER:\s*)(\d{6})", re.IGNORECASE)

# Professional Fees: $12,500 (same line)
PRO_FEES_RE = re.compile(
    r"(Professional Fees:\s*)(\$?\s*[\d,]+(?:\.\d{2})?)",
    re.IGNORECASE
)

# Title line hint
TITLE_HINT_RE = re.compile(r"\bSOC Audit\b.*\bInvoice\b", re.IGNORECASE)

# Mgmt Rep Letter patterns
LONG_DATE_RE = re.compile(r"(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2})")
STANDALONE_LONG_DATE_RE = re.compile(
    r"^\s*((" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2}))\s*$"
)

# Review period example: "February 1, 2023 to January 31, 2024"
REVIEW_PERIOD_RE = re.compile(
    r"("
    r"(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2})"
    r"\s+to\s+"
    r"(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2})"
    r")"
)

# "as of March 20, 2024" (we replace only the date portion)
AS_OF_DATE_RE = re.compile(
    r"(as of\s+)((" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2}))",
    re.IGNORECASE
)


# -----------------------------
# General helpers
# -----------------------------
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def _is_zip(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() == "zip"


def _parse_rollforward_date(iso_yyyy_mm_dd: str) -> date:
    return datetime.strptime(iso_yyyy_mm_dd, "%Y-%m-%d").date()


def _format_long_date(d: date) -> str:
    # "January 15, 2025" (remove leading zero day)
    return d.strftime("%B %d, %Y").replace(" 0", " ")


def _add_one_month_same_day(d: date) -> date:
    year = d.year
    month = d.month + 1
    if month == 13:
        month = 1
        year += 1

    last_day = calendar.monthrange(year, month)[1]
    day = min(d.day, last_day)
    return date(year, month, day)


def _format_currency(user_text: str) -> str:
    s = (user_text or "").strip()
    if not s:
        raise ValueError("New amount is empty.")

    cleaned = s.replace("$", "").replace(" ", "")
    if not re.fullmatch(r"[\d,]+(\.\d{2})?", cleaned):
        raise ValueError("Amount must look like 12500, 12,500, $12,500, or 12500.00")

    has_cents = re.search(r"\.\d{2}$", cleaned) is not None
    numeric = float(cleaned.replace(",", ""))

    if has_cents:
        return f"${numeric:,.2f}"
    return f"${numeric:,.0f}"


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def build_initial_output_basename_from_uploaded_filename(uploaded_filename_raw: str, rollforward_year: int) -> str:
    """
    Initial invoice output naming:
      "[Client Name] - 202X SOC Audit - Invoice 1"
    Uses original uploaded filename for the base, but rolls forward the year and forces Invoice 1.
    """
    base = os.path.splitext(os.path.basename(uploaded_filename_raw))[0]
    base = base.replace("_", " ")
    base = re.sub(r"\s*-\s*", " - ", base)
    base = re.sub(r"\s+", " ", base).strip()

    year_pat = re.compile(r"(?<!\d)20\d{2}(?!\d)")
    if year_pat.search(base):
        base = year_pat.sub(str(rollforward_year), base, count=1)
    else:
        base = f"{base} - {rollforward_year}"

    invoice_pat = re.compile(r"(Invoice)[\s_]+(\d+)", re.IGNORECASE)
    if invoice_pat.search(base):
        base = invoice_pat.sub("Invoice 1", base, count=1)
    else:
        base = f"{base} - Invoice 1"

    base = re.sub(r"\s+", " ", base).strip()
    return base


def build_final_output_basename_from_uploaded_filename(uploaded_filename_raw: str) -> str:
    """
    Final invoice output naming:
      Keep the ORIGINAL uploaded filename base, but change "Invoice 1" -> "Invoice 2".
      Do NOT roll-forward the year in the filename (per your rule #5).
    """
    base = os.path.splitext(os.path.basename(uploaded_filename_raw))[0]
    base = base.replace("_", " ")
    base = re.sub(r"\s*-\s*", " - ", base)
    base = re.sub(r"\s+", " ", base).strip()

    invoice_pat = re.compile(r"(Invoice)[\s_]+(\d+)", re.IGNORECASE)
    if invoice_pat.search(base):
        base = invoice_pat.sub("Invoice 2", base, count=1)
    else:
        base = f"{base} - Invoice 2"

    base = re.sub(r"\s+", " ", base).strip()
    return base


def build_mgmt_rep_output_filename(raw_uploaded_filename: str) -> str:
    """
    Mgmt Rep output naming rule:
      "Change nothing about the file name except +1 to the year included in the originally uploaded file name."
    Example:
      "Management Rep Letter - 2024 Temporary Housing Directory.docx" -> "... 2025 ... .docx"
    We treat the FIRST 20xx token as the year to increment.
    """
    original_base = os.path.splitext(os.path.basename(raw_uploaded_filename))[0]
    ext = ".docx"  # output is always docx

    year_pat = re.compile(r"(?<!\d)(20\d{2})(?!\d)")
    m = year_pat.search(original_base)
    if not m:
        # If no year is present, do not invent one — keep filename identical and just force .docx
        return f"{original_base}{ext}"

    old_year = int(m.group(1))
    new_year = old_year + 1
    new_base = year_pat.sub(str(new_year), original_base, count=1)

    return f"{new_base}{ext}"


def _sanitize_filename_for_windows(name: str) -> str:
    safe = re.sub(r'[<>:"/\\|?*\n\r\t]', " ", name).strip()
    safe = re.sub(r"\s+", " ", safe)
    return safe


# -----------------------------
# Run-preserving text editing
# -----------------------------
def _get_runs_text(paragraph) -> str:
    return "".join(r.text for r in paragraph.runs)


def _build_run_spans(paragraph) -> List[Tuple[int, int]]:
    spans = []
    idx = 0
    for r in paragraph.runs:
        start = idx
        idx += len(r.text or "")
        end = idx
        spans.append((start, end))
    return spans


def _find_run_index_at_char(paragraph, char_index: int) -> Optional[int]:
    spans = _build_run_spans(paragraph)
    for i, (s, e) in enumerate(spans):
        if s <= char_index < e:
            return i
    if spans and char_index == spans[-1][1]:
        return len(spans) - 1
    return None


def _replace_span_in_runs(paragraph, start: int, end: int, replacement: str) -> bool:
    if start >= end:
        return False

    runs = paragraph.runs
    if not runs:
        return False

    spans = _build_run_spans(paragraph)

    start_run = None
    end_run = None
    for i, (s, e) in enumerate(spans):
        if start_run is None and s <= start < e:
            start_run = i
        if s < end <= e:
            end_run = i
            break

    if start_run is None:
        return False
    if end_run is None:
        end_run = len(runs) - 1

    s_start, _ = spans[start_run]
    e_start, _ = spans[end_run]
    start_off = start - s_start
    end_off = end - e_start

    before = runs[start_run].text[:start_off]
    after = runs[end_run].text[end_off:]

    runs[start_run].text = before + replacement + after

    for j in range(start_run + 1, end_run + 1):
        runs[j].text = ""

    return True


def _replace_whole_word_in_runs(paragraph, pattern: re.Pattern, repl_func) -> bool:
    changed = False
    for r in paragraph.runs:
        if not r.text:
            continue
        new_text = pattern.sub(repl_func, r.text)
        if new_text != r.text:
            r.text = new_text
            changed = True
    return changed


def _preserve_case(match: re.Match, lower: str, title: str, upper: str) -> str:
    w = match.group(0)
    if w.isupper():
        return upper
    if w[:1].isupper():
        return title
    return lower


def _preserve_case_final_to_initial(match: re.Match) -> str:
    return _preserve_case(match, lower="initial", title="Initial", upper="INITIAL")


def _preserve_case_initial_to_final(match: re.Match) -> str:
    return _preserve_case(match, lower="final", title="Final", upper="FINAL")


def _find_label_literal(full_text: str, label_regex: re.Pattern) -> Optional[str]:
    """
    Find the label portion in full_text and return it exactly as it appears (preserves casing),
    e.g., "Due Date:".
    """
    m = label_regex.search(full_text)
    if not m:
        return None
    return m.group(0)


def _replace_value_after_label_preserve_bold(paragraph, label_literal: str, new_value: str) -> bool:
    """
    Replace everything after label_literal (e.g., "DATE:" or "Due Date:") with new_value,
    preserving bold label / non-bold value styling.
    """
    full = _get_runs_text(paragraph)
    label_pos = full.find(label_literal)
    if label_pos == -1:
        return False

    start = label_pos + len(label_literal)
    while start < len(full) and full[start] == " ":
        start += 1
    end = len(full)

    label_run_idx = _find_run_index_at_char(paragraph, label_pos)
    if label_run_idx is None:
        return False

    value_run_idx = None
    for i in range(label_run_idx + 1, len(paragraph.runs)):
        run = paragraph.runs[i]
        if run.bold is False or run.bold is None:
            value_run_idx = i
            break

    if value_run_idx is None:
        return _replace_span_in_runs(paragraph, start, end, new_value)

    spans = _build_run_spans(paragraph)
    vs, _ = spans[value_run_idx]
    replace_start = max(start, vs)
    return _replace_span_in_runs(paragraph, replace_start, end, new_value)


def _replace_by_regex_on_full_text_preserve_runs(paragraph, regex: re.Pattern, repl_func) -> bool:
    """
    If regex has capturing groups, this function replaces span(1).
    If regex has no groups, it replaces the full match span.
    """
    full = _get_runs_text(paragraph)
    matches = list(regex.finditer(full))
    if not matches:
        return False

    changed = False
    for m in reversed(matches):
        start, end = m.span(1) if m.lastindex else m.span()
        replacement = repl_func(m)
        if _replace_span_in_runs(paragraph, start, end, replacement):
            changed = True

    return changed


def _replace_by_regex_group_preserve_runs(paragraph, regex: re.Pattern, group_index: int, repl_func) -> bool:
    """
    Replace a specific capturing group span (e.g., group 2) while preserving runs.
    Useful when you want to keep surrounding text like "as of ".
    """
    full = _get_runs_text(paragraph)
    matches = list(regex.finditer(full))
    if not matches:
        return False

    changed = False
    for m in reversed(matches):
        start, end = m.span(group_index)
        replacement = repl_func(m)
        if _replace_span_in_runs(paragraph, start, end, replacement):
            changed = True

    return changed


# -----------------------------
# LibreOffice conversion
# -----------------------------
def _get_soffice_cmd() -> List[str]:
    env_path = os.environ.get("LIBREOFFICE_PATH", "").strip()
    if env_path:
        return [env_path]

    soffice = shutil.which("soffice")
    if soffice:
        return [soffice]

    win_default = r"C:\Program Files\LibreOffice\program\soffice.exe"
    if os.path.exists(win_default):
        return [win_default]

    raise RuntimeError(
        "LibreOffice not found. Install LibreOffice or set LIBREOFFICE_PATH "
        "to the full path of soffice.exe/soffice."
    )


def convert_to_docx(input_path: str, out_dir: str) -> str:
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".docx":
        return input_path
    if ext != ".doc":
        raise ValueError("convert_to_docx only accepts .doc or .docx")

    cmd = _get_soffice_cmd() + [
        "--headless", "--nologo", "--nolockcheck", "--nodefault", "--norestore",
        "--convert-to", "docx",
        "--outdir", out_dir,
        input_path
    ]

    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"DOC→DOCX conversion failed:\n{res.stderr or res.stdout}")

    base = os.path.splitext(os.path.basename(input_path))[0]
    converted = os.path.join(out_dir, base + ".docx")
    if not os.path.exists(converted) or os.path.getsize(converted) == 0:
        raise RuntimeError("DOC→DOCX conversion did not produce a valid .docx output.")
    return converted


def convert_docx_to_pdf(input_docx_path: str, out_dir: str) -> str:
    if os.path.splitext(input_docx_path)[1].lower() != ".docx":
        raise ValueError("convert_docx_to_pdf requires a .docx input")

    cmd = _get_soffice_cmd() + [
        "--headless", "--nologo", "--nolockcheck", "--nodefault", "--norestore",
        "--convert-to", "pdf",
        "--outdir", out_dir,
        input_docx_path
    ]

    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"DOCX→PDF conversion failed:\n{res.stderr or res.stdout}")

    base = os.path.splitext(os.path.basename(input_docx_path))[0]
    pdf_path = os.path.join(out_dir, base + ".pdf")
    if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
        raise RuntimeError("DOCX→PDF conversion did not produce a valid .pdf output.")
    return pdf_path


# -----------------------------
# Rollforward Rules (Initial Invoice)
# -----------------------------
def apply_initial_invoice_rules(
    doc: Document,
    rollforward_dt: date,
    amount_mode: str,
    new_amount_text: Optional[str]
) -> None:
    rf_date_str = _format_long_date(rollforward_dt)
    due_dt = _add_one_month_same_day(rollforward_dt)
    due_date_str = _format_long_date(due_dt)
    target_year = rollforward_dt.year

    new_amount_formatted = None
    if amount_mode == "changed":
        new_amount_formatted = _format_currency(new_amount_text or "")

    # Update title inside the document to Invoice 1
    for p in _iter_all_paragraphs(doc):
        if TITLE_HINT_RE.search(_get_runs_text(p)):
            full = _get_runs_text(p)
            parts = full.split(" - ", 1)
            if len(parts) == 2:
                client_name = parts[0].strip()
                new_title = f"{client_name} - {target_year} SOC Audit - Invoice 1"
                if p.runs:
                    p.runs[0].text = new_title
                    for r in p.runs[1:]:
                        r.text = ""
            break

    for p in _iter_all_paragraphs(doc):
        full = _get_runs_text(p)
        if not full.strip():
            continue

        # DATE:
        if DATE_LINE_RE.match(full.strip()):
            _replace_value_after_label_preserve_bold(p, "DATE:", rf_date_str)
            continue

        # Due Date:
        if DUE_DATE_LINE_RE.match(full.strip()):
            label_literal = _find_label_literal(full, re.compile(r"due date:", re.IGNORECASE))
            if label_literal:
                _replace_value_after_label_preserve_bold(p, label_literal, due_date_str)
            continue

        # final -> initial (word only)
        _replace_whole_word_in_runs(p, FINAL_WORD_RE, _preserve_case_final_to_initial)

        # year before "SOC Audit engagement" +1
        _replace_by_regex_on_full_text_preserve_runs(
            p,
            ENGAGEMENT_YEAR_RE,
            lambda m: str(int(m.group(1)) + 1)
        )

        # Professional Fees amount (optional)
        if new_amount_formatted:
            full2 = _get_runs_text(p)
            m = PRO_FEES_RE.search(full2)
            if m:
                s_start, s_end = m.span(2)
                _replace_span_in_runs(p, s_start, s_end, new_amount_formatted)

        # Invoice number update (digits 3-4 >= 10): +2 and last two "2" + random(2/4/6/8)
        full3 = _get_runs_text(p)
        matches = list(INVOICE_NUMBER_RE.finditer(full3))
        if matches:
            for mm in reversed(matches):
                six = mm.group(2)
                two_digit = int(six[2:4])
                if two_digit < 10:
                    continue

                new_two = (two_digit + 2) % 100
                new_two_str = f"{new_two:02d}"

                candidates = ["2", "4", "6", "8"]
                pick = candidates[(uuid.uuid4().int % len(candidates))]
                last_two = "2" + pick

                new_six = six[:2] + new_two_str + last_two

                s_start, s_end = mm.span(2)
                _replace_span_in_runs(p, s_start, s_end, new_six)


# -----------------------------
# Rollforward Rules (Final Invoice)
# -----------------------------
def apply_final_invoice_rules(doc: Document, rollforward_dt: date) -> None:
    """
    Final invoice rules:
      1) Invoice number: last 2 digits become "8" + random(2/4/6/8)
      2) DATE: set to calendar selection
      3) Replace "initial" -> "final" (whole word), preserve case
      4) Due date: one month after selected date, same day
      5) Preserve formatting by editing runs (no paragraph.text assignments)
    """
    rf_date_str = _format_long_date(rollforward_dt)
    due_dt = _add_one_month_same_day(rollforward_dt)
    due_date_str = _format_long_date(due_dt)

    # (Optional) Update title inside the document to Invoice 2
    for p in _iter_all_paragraphs(doc):
        if TITLE_HINT_RE.search(_get_runs_text(p)):
            full = _get_runs_text(p)
            parts = full.split(" - ", 1)
            if len(parts) == 2:
                client_name = parts[0].strip()
                m_year = re.search(r"(?<!\d)(20\d{2})(?!\d)", full)
                year_in_title = m_year.group(1) if m_year else str(rollforward_dt.year)
                new_title = f"{client_name} - {year_in_title} SOC Audit - Invoice 2"
                if p.runs:
                    p.runs[0].text = new_title
                    for r in p.runs[1:]:
                        r.text = ""
            break

    for p in _iter_all_paragraphs(doc):
        full = _get_runs_text(p)
        if not full.strip():
            continue

        # Rule #2: DATE:
        if DATE_LINE_RE.match(full.strip()):
            _replace_value_after_label_preserve_bold(p, "DATE:", rf_date_str)
            continue

        # Rule #4: Due Date:
        if DUE_DATE_LINE_RE.match(full.strip()):
            label_literal = _find_label_literal(full, re.compile(r"due date:", re.IGNORECASE))
            if label_literal:
                _replace_value_after_label_preserve_bold(p, label_literal, due_date_str)
            continue

        # Rule #3: initial -> final
        _replace_whole_word_in_runs(p, INITIAL_WORD_RE, _preserve_case_initial_to_final)

        # Rule #1: Invoice number last 2 digits -> "8" + random(2/4/6/8)
        full3 = _get_runs_text(p)
        matches = list(INVOICE_NUMBER_RE.finditer(full3))
        if matches:
            for mm in reversed(matches):
                six = mm.group(2)

                candidates = ["2", "4", "6", "8"]
                pick = candidates[(uuid.uuid4().int % len(candidates))]
                last_two = "8" + pick

                new_six = six[:4] + last_two

                s_start, s_end = mm.span(2)
                _replace_span_in_runs(p, s_start, s_end, new_six)


# -----------------------------
# Rollforward Rules (Management Rep Letter)
# -----------------------------
def apply_mgmt_rep_letter_rules(doc: Document, selected_dt: date) -> None:
    """
    Mgmt Rep Letter rules:
      1) Replace top-of-letter date (standalone "Month Day, Year") with selected date
         AND replace date portion after "as of " with selected date.
      2) Roll forward ALL review periods "Month Day, YYYY to Month Day, YYYY" by +1 year.
      3) Preserve formatting by editing runs only (no paragraph.text assignments).
    """
    selected_str = _format_long_date(selected_dt)

    for p in _iter_all_paragraphs(doc):
        full = _get_runs_text(p)

        if not full or not full.strip():
            continue

        # Rule #1a: Standalone date line (top of letter style)
        m_standalone = STANDALONE_LONG_DATE_RE.match(full)
        if m_standalone:
            # Replace only the date span (group 1 is the full date)
            date_text = m_standalone.group(1)
            idx = full.find(date_text)
            if idx != -1:
                _replace_span_in_runs(p, idx, idx + len(date_text), selected_str)
            continue

        # Rule #1b: "as of <DATE>" (replace only the date group)
        _replace_by_regex_group_preserve_runs(
            p,
            AS_OF_DATE_RE,
            group_index=2,
            repl_func=lambda m: selected_str
        )

        # Rule #2: Review periods by +1 year (replace the whole match (group 1))
        def _roll_period(m: re.Match) -> str:
            # groups: 2 month1, 3 day1, 4 year1, 5 month2, 6 day2, 7 year2
            month1 = m.group(2)
            day1 = m.group(3)
            year1 = int(m.group(4)) + 1
            month2 = m.group(5)
            day2 = m.group(6)
            year2 = int(m.group(7)) + 1
            return f"{month1} {int(day1)}, {year1} to {month2} {int(day2)}, {year2}"

        _replace_by_regex_on_full_text_preserve_runs(p, REVIEW_PERIOD_RE, _roll_period)


# ============================================================
# NEW FEATURE: Information Requests Received Review
# ============================================================

# Numeric Control Point like 1.2, 5.10, 7.4 etc.
# IMPORTANT: negative lookbehind blocks matching inside SOC 2 refs like "CC6.4"
INFOREQ_CP_REGEX = re.compile(r"(?<![A-Za-z])(\d+(?:\.\d+)+)")
PROVIDE_BY_MON_DD_RE = re.compile(
    r"\bprovide\s+by\s+([A-Z]{3})\s+(\d{2})\b",
    re.IGNORECASE
)


def _cell_text(cell) -> str:
    # Join paragraphs, then normalize "weird Word whitespace"
    txt = "\n".join([p.text for p in cell.paragraphs])

    # Common Word gremlins:
    txt = txt.replace("\xa0", " ")      # non-breaking space
    txt = txt.replace("\u200b", "")     # zero-width space
    txt = txt.replace("\u2011", "-")    # non-breaking hyphen
    txt = txt.replace("\u2013", "-")    # en dash
    txt = txt.replace("\u2014", "-")    # em dash

    # collapse all whitespace to single spaces
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt


def _find_header_row_and_cols(table):
    """
    Detects header row and identifies:
      - Control Point column
      - Request column (either 'Test Results...' OR 'Information to be Provided by...')
    """

    def norm(s):
        s = (s or "").replace("\xa0", " ").replace("\u200b", "")
        s = re.sub(r"\s+", " ", s).strip().lower()
        return s

    scan_rows = min(10, len(table.rows))
    for r_idx in range(scan_rows):
        row = table.rows[r_idx]
        texts = [norm(_cell_text(c)) for c in row.cells]

        cp_col = None
        req_col = None

        for c_idx, t in enumerate(texts):

            # CONTROL POINT detection
            if "control" in t and "point" in t:
                cp_col = c_idx

            # EZ FORM column
            if "test" in t and "result" in t:
                req_col = c_idx

            # SOC 1 column
            if "information" in t and "provided" in t:
                req_col = c_idx

        if cp_col is not None and req_col is not None:
            return r_idx, cp_col, req_col

    return None, None, None

def _extract_requested_control_points_from_docx(docx_path: str) -> List[Dict]:
    doc = Document(docx_path)
    requested = []

    for t_idx, table in enumerate(doc.tables):
        hdr_r, cp_col, tr_col = _find_header_row_and_cols(table)
        if hdr_r is None:
            continue

        last_cp = None  # carry-down for blank CP rows

        for r_idx in range(hdr_r + 1, len(table.rows)):
            row = table.rows[r_idx]

            try:
                cp_raw = _cell_text(row.cells[cp_col])
                tr_text = _cell_text(row.cells[tr_col])
            except Exception:
                continue

            # Update last_cp when we see a numeric control point
            m_cp = INFOREQ_CP_REGEX.search(cp_raw or "")
            if m_cp:
                last_cp = m_cp.group(1)

            # Only count as a request if it matches your strict format
            if not PROVIDE_BY_MON_DD_RE.search(tr_text or ""):
                continue

            # Use this row's CP or the carry-down CP
            control_point = m_cp.group(1) if m_cp else last_cp
            if not control_point:
                continue

            requested.append({
                "control_point": control_point,
                "request_text": tr_text.strip(),
                "table_index": t_idx,
                "row_index": r_idx
            })

    # Deduplicate by control point (combine request text if needed)
    seen = {}
    for it in requested:
        cp = it["control_point"]
        if cp not in seen:
            seen[cp] = it
        else:
            if it["request_text"] and it["request_text"] not in (seen[cp].get("request_text") or ""):
                seen[cp]["request_text"] = (seen[cp]["request_text"] + " | " + it["request_text"]).strip(" |")

    return list(seen.values())


def _extract_zip(zip_path: str, extract_dir: str) -> None:
    os.makedirs(extract_dir, exist_ok=True)
    with zipfile.ZipFile(zip_path, "r") as z:
        z.extractall(extract_dir)


def _index_evidence(extract_dir: str) -> Tuple[List[Dict], Dict[str, List[str]]]:
    """
    Builds:
      - evidence_index: list of file entries + CP tokens found
      - cp_to_files: CP token -> list of relpaths that contain that CP token
    We scan the RELATIVE PATH (folders + filename), so nested folder naming also counts.
    """
    evidence_index = []
    cp_to_files: Dict[str, List[str]] = {}

    for root, _, files in os.walk(extract_dir):
        for fn in files:
            full_path = os.path.join(root, fn)
            rel_path = os.path.relpath(full_path, extract_dir).replace("\\", "/")

            haystack = rel_path.lower()

            # Find all numeric CP tokens NOT preceded by letters (ignores cc6.4)
            cps = [m.group(1) for m in INFOREQ_CP_REGEX.finditer(haystack)]
            cps_set = sorted(set(cps))

            evidence_index.append({
                "rel_path": rel_path,
                "filename": fn,
                "control_points_found": cps_set
            })

            for cp in cps_set:
                cp_to_files.setdefault(cp, []).append(rel_path)

    for cp in list(cp_to_files.keys()):
        cp_to_files[cp] = sorted(set(cp_to_files[cp]))

    return evidence_index, cp_to_files


def _write_xlsx_report(
    out_xlsx_path: str,
    requested_items: List[Dict],
    evidence_index: List[Dict]
) -> None:
    if openpyxl is None:
        raise RuntimeError("openpyxl not installed. Install openpyxl or switch to CSV output.")

    wb = Workbook()
    ws_summary = wb.active
    ws_summary.title = "Summary"

    total = len(requested_items)
    found = sum(1 for it in requested_items if it.get("evidence_found"))
    missing = total - found
    # -----------------------------
    # Missing Control Points list
    # -----------------------------
    missing_cps = sorted([
        it["control_point"]
        for it in requested_items
        if not it.get("evidence_found")
    ])

    ws_summary["A8"] = "Missing Control Points"
    ws_summary["A9"] = "(Evidence not found in ZIP filenames)"

    if not missing_cps:
        ws_summary["A11"] = "None ✅"
    else:
        ws_summary["A11"] = "Control Point"
        start_row = 12

        for i, cp in enumerate(missing_cps):
            ws_summary[f"A{start_row + i}"] = cp
    ws_summary["A1"] = "Information Requests Received Review"
    ws_summary["A2"] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ws_summary["A4"] = "Total Requested Control Points (Provide by)"
    ws_summary["B4"] = total
    ws_summary["A5"] = "Evidence Found"
    ws_summary["B5"] = found
    ws_summary["A6"] = "Missing"
    ws_summary["B6"] = missing

    ws_req = wb.create_sheet("Requests")
    ws_req.append([
        "Control Point",
        "Requested? (Provide by)",
        "Evidence Found?",
        "Matching Evidence Files",
        "Request Text (Test Results)",
        "Table Index",
        "Row Index"
    ])

    for it in requested_items:
        ws_req.append([
            it.get("control_point", ""),
            "Yes",
            "Yes" if it.get("evidence_found") else "No",
            ", ".join(it.get("matching_files", [])) if it.get("matching_files") else "",
            it.get("request_text", ""),
            it.get("table_index", ""),
            it.get("row_index", "")
        ])

    ws_ev = wb.create_sheet("Evidence_Index")
    ws_ev.append(["Relative Path", "Filename", "Control Points Found (numeric)"])
    for ev in evidence_index:
        ws_ev.append([ev["rel_path"], ev["filename"], ", ".join(ev["control_points_found"])])

    # Basic sizing
    for ws in (ws_summary, ws_req, ws_ev):
        for col in range(1, ws.max_column + 1):
            letter = get_column_letter(col)
            ws.column_dimensions[letter].width = 35

    wb.save(out_xlsx_path)


def _write_csv_report(
    out_csv_path: str,
    requested_items: List[Dict],
    evidence_index: List[Dict]
) -> None:
    import csv

    # One CSV with two sections (simple fallback)
    with open(out_csv_path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["Information Requests Received Review"])
        w.writerow(["Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
        w.writerow([])
        w.writerow(["Requests"])
        w.writerow([
            "Control Point",
            "Requested? (Provide by)",
            "Evidence Found?",
            "Matching Evidence Files",
            "Request Text (Test Results)",
            "Table Index",
            "Row Index"
        ])
        for it in requested_items:
            w.writerow([
                it.get("control_point", ""),
                "Yes",
                "Yes" if it.get("evidence_found") else "No",
                " | ".join(it.get("matching_files", [])) if it.get("matching_files") else "",
                it.get("request_text", ""),
                it.get("table_index", ""),
                it.get("row_index", "")
            ])
        w.writerow([])
        w.writerow(["Evidence Index"])
        w.writerow(["Relative Path", "Filename", "Control Points Found (numeric)"])
        for ev in evidence_index:
            w.writerow([ev["rel_path"], ev["filename"], " | ".join(ev["control_points_found"])])


# -----------------------------
# Routes
# -----------------------------
@app.route("/", methods=["GET"])
def home():
    return render_template("home.html")


@app.route("/invoice", methods=["GET"])
def invoice_page():
    return render_template("invoice.html")


@app.route("/mgmt-rep", methods=["GET"])
def mgmt_rep_page():
    return render_template("mgmt_rep.html")


# NEW page route (template you'll add below)
@app.route("/info-requests-received", methods=["GET"])
def info_requests_received_page():
    return render_template("info_requests_received.html")


@app.route("/rollforward", methods=["POST"])
def rollforward():
    invoice_type = request.form.get("invoice_type", "initial").strip().lower()
    rollforward_date_raw = (request.form.get("rollforward_date") or "").strip()

    # Only used for Initial invoice
    amount_mode = (request.form.get("amount_mode") or "same").strip().lower()
    new_amount = (request.form.get("new_amount") or "").strip()

    file = request.files.get("invoice_file")

    if not file or file.filename.strip() == "":
        return "No file uploaded.", 400
    if not allowed_file(file.filename):
        return f"File type not allowed. Allowed: {sorted(ALLOWED_EXTENSIONS)}", 400
    if not rollforward_date_raw:
        return "Rollforward Date is required.", 400

    try:
        rollforward_dt = _parse_rollforward_date(rollforward_date_raw)
    except Exception:
        return "Invalid Rollforward Date format.", 400

    # Keep raw filename for output naming (preserves spaces), but use secure for saving
    raw_uploaded_name = file.filename
    safe_uploaded_name = secure_filename(file.filename)

    upload_id = uuid.uuid4().hex[:10]
    uploaded_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{upload_id}__{safe_uploaded_name}")
    file.save(uploaded_path)

    try:
        # Convert to DOCX if needed
        processing_docx_path = convert_to_docx(uploaded_path, app.config["UPLOAD_FOLDER"])
        doc = Document(processing_docx_path)

        if invoice_type == "initial":
            if amount_mode not in {"same", "changed"}:
                return "Invalid Invoice Amount selection.", 400
            if amount_mode == "changed" and not new_amount:
                return "New Amount is required when Invoice Amount = Has changed.", 400

            apply_initial_invoice_rules(
                doc=doc,
                rollforward_dt=rollforward_dt,
                amount_mode=amount_mode,
                new_amount_text=new_amount if amount_mode == "changed" else None
            )

            # Output name rolls forward year and forces Invoice 1
            out_base = build_initial_output_basename_from_uploaded_filename(raw_uploaded_name, rollforward_dt.year)
            safe_base = _sanitize_filename_for_windows(out_base)

            # Save DOCX
            out_docx_path = os.path.join(app.config["OUTPUT_FOLDER"], f"{safe_base}.docx")
            doc.save(out_docx_path)

            # Convert DOCX -> PDF
            out_pdf_path = convert_docx_to_pdf(out_docx_path, app.config["OUTPUT_FOLDER"])

            # Zip both
            zip_name = f"{safe_base} (DOCX+PDF).zip"
            zip_path = os.path.join(app.config["OUTPUT_FOLDER"], zip_name)

            with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
                z.write(out_docx_path, arcname=os.path.basename(out_docx_path))
                z.write(out_pdf_path, arcname=os.path.basename(out_pdf_path))

            return send_file(zip_path, as_attachment=True, download_name=zip_name)

        elif invoice_type == "final":
            apply_final_invoice_rules(doc, rollforward_dt)

            # Output name keeps original, only Invoice 1 -> Invoice 2
            out_base = build_final_output_basename_from_uploaded_filename(raw_uploaded_name)
            safe_base = _sanitize_filename_for_windows(out_base)

            out_docx_path = os.path.join(app.config["OUTPUT_FOLDER"], f"{safe_base}.docx")
            doc.save(out_docx_path)

            # Rule #5: DOCX only
            return send_file(out_docx_path, as_attachment=True, download_name=os.path.basename(out_docx_path))

        else:
            return "Invalid invoice_type. Use 'initial' or 'final'.", 400

    except Exception as e:
        return f"Rollforward failed: {str(e)}", 500


@app.route("/rollforward/mgmt-rep", methods=["POST"])
def rollforward_mgmt_rep():
    rollforward_date_raw = (request.form.get("rollforward_date") or "").strip()
    file = request.files.get("rep_file")

    if not file or file.filename.strip() == "":
        return "No file uploaded.", 400
    if not allowed_file(file.filename):
        return f"File type not allowed. Allowed: {sorted(ALLOWED_EXTENSIONS)}", 400
    if not rollforward_date_raw:
        return "Rollforward Date is required.", 400

    try:
        selected_dt = _parse_rollforward_date(rollforward_date_raw)
    except Exception:
        return "Invalid Rollforward Date format.", 400

    raw_uploaded_name = file.filename
    safe_uploaded_name = secure_filename(file.filename)

    upload_id = uuid.uuid4().hex[:10]
    uploaded_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{upload_id}__{safe_uploaded_name}")
    file.save(uploaded_path)

    try:
        # Convert to DOCX if needed
        processing_docx_path = convert_to_docx(uploaded_path, app.config["UPLOAD_FOLDER"])
        doc = Document(processing_docx_path)

        # Apply mgmt rep rules
        apply_mgmt_rep_letter_rules(doc, selected_dt)

        # Output filename: same as uploaded except first year +1, docx only
        out_filename = build_mgmt_rep_output_filename(raw_uploaded_name)
        out_filename_safe = _sanitize_filename_for_windows(out_filename)

        out_docx_path = os.path.join(app.config["OUTPUT_FOLDER"], out_filename_safe)
        doc.save(out_docx_path)

        return send_file(out_docx_path, as_attachment=True, download_name=os.path.basename(out_docx_path))

    except Exception as e:
        return f"Mgmt Rep rollforward failed: {str(e)}", 500


# NEW: run review
# NEW: run review
@app.route("/info-requests-received/run", methods=["POST"])
def info_requests_received_run():
    req_file = request.files.get("requests_docx")
    zip_file = request.files.get("evidence_zip")

    if not req_file or req_file.filename.strip() == "":
        return "No Information Requests file uploaded.", 400
    if not zip_file or zip_file.filename.strip() == "":
        return "No evidence ZIP uploaded.", 400

    if not allowed_file(req_file.filename):
        return f"Information Requests file type not allowed. Allowed: {sorted(ALLOWED_EXTENSIONS)}", 400
    if not _is_zip(zip_file.filename):
        return "Evidence must be a .zip file.", 400

    upload_id = uuid.uuid4().hex[:10]
    safe_req_name = secure_filename(req_file.filename)
    safe_zip_name = secure_filename(zip_file.filename)

    req_uploaded_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{upload_id}__{safe_req_name}")
    zip_uploaded_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{upload_id}__{safe_zip_name}")

    req_file.save(req_uploaded_path)
    zip_file.save(zip_uploaded_path)

    extract_dir = os.path.join(app.config["UPLOAD_FOLDER"], f"{upload_id}__evidence_extracted")
    os.makedirs(extract_dir, exist_ok=True)

    try:
        # Convert doc -> docx if needed
        processing_docx_path = convert_to_docx(req_uploaded_path, app.config["UPLOAD_FOLDER"])

        # 1) parse requested CPs (Provide by rows)
        requested_items = _extract_requested_control_points_from_docx(processing_docx_path)

        # 2) extract + index evidence
        _extract_zip(zip_uploaded_path, extract_dir)
        evidence_index, cp_to_files = _index_evidence(extract_dir)

        # 3) match
        for it in requested_items:
            cp = it["control_point"]
            matches = cp_to_files.get(cp, [])
            it["matching_files"] = matches
            it["evidence_found"] = bool(matches)

        # -----------------------------
        # Excel Output + UI Summary
        # -----------------------------
        out_base = f"Info Requests Received Review - {upload_id}"

        if openpyxl is None:
            return "openpyxl is required to generate the Excel output. Please run: pip install openpyxl", 500

        out_xlsx = os.path.join(
            app.config["OUTPUT_FOLDER"],
            f"{out_base}.xlsx"
        )

        _write_xlsx_report(
            out_xlsx,
            requested_items,
            evidence_index
        )

        # Build missing CP list for UI
        missing_cps = sorted([
            it["control_point"]
            for it in requested_items
            if not it.get("evidence_found")
        ])

        # If there are 0 requests found, we do NOT want to show "100% provided"
        provided_all = (len(requested_items) > 0 and len(missing_cps) == 0)

        download_url = f"/info-requests-received/download/{upload_id}"

        # Render the same page with results.
        # The template will auto-trigger the Excel download via JS.
        return render_template(
            "info_requests_received.html",
            results_ready=True,
            provided_all=provided_all,
            missing_cps=missing_cps,
            total_requested=len(requested_items),
            total_missing=len(missing_cps),
            download_url=download_url
        )

    except Exception as e:
        return f"Info Requests Received Review failed: {str(e)}", 500

    finally:
        # Leave uploads for troubleshooting; if you want auto-cleanup later we can turn it on.
        pass

# NEW: download the generated Excel
@app.route("/info-requests-received/download/<upload_id>", methods=["GET"])
def info_requests_received_download(upload_id):
    filename = f"Info Requests Received Review - {upload_id}.xlsx"
    path = os.path.join(app.config["OUTPUT_FOLDER"], filename)

    if not os.path.exists(path):
        return "File not found. Please run the review again.", 404

    return send_file(
        path,
        as_attachment=True,
        download_name=filename
    )

if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)