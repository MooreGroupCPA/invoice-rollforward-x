import os
import re
import uuid
import zipfile
import shutil
import subprocess
import calendar
from datetime import datetime, date
from typing import Optional, List, Tuple

from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename

from docx import Document


APP_DIR = os.path.dirname(os.path.abspath(__file__))

UPLOAD_FOLDER = os.path.join(APP_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(APP_DIR, "output")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {"doc", "docx"}

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 30 * 1024 * 1024  # 30 MB


# -----------------------------
# Regex / Patterns
# -----------------------------
MONTHS = (
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December"
)

DATE_LINE_RE = re.compile(
    r"^DATE:\s+(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(\d{4})\s*$"
)

# Matches "Due Date:" or "DUE DATE:" etc.
DUE_DATE_LINE_RE = re.compile(
    r"^(due date):\s+(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(\d{4})\s*$",
    re.IGNORECASE
)

ENGAGEMENT_YEAR_RE = re.compile(r"\b(20\d{2})(?=\s+SOC Audit engagement\b)")

# Whole-word replacements
FINAL_WORD_RE = re.compile(r"\bfinal\b", re.IGNORECASE)
INITIAL_WORD_RE = re.compile(r"\binitial\b", re.IGNORECASE)

# INVOICE NUMBER: 6 digits
INVOICE_NUMBER_RE = re.compile(r"(INVOICE NUMBER:\s*)(\d{6})", re.IGNORECASE)

# Professional Fees: $12,500 (same line)
PRO_FEES_RE = re.compile(
    r"(Professional Fees:\s*)(\$?\s*[\d,]+(?:\.\d{2})?)",
    re.IGNORECASE
)

# Title line hint
TITLE_HINT_RE = re.compile(r"\bSOC Audit\b.*\bInvoice\b", re.IGNORECASE)

# Mgmt Rep Letter patterns
LONG_DATE_RE = re.compile(r"(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2})")
STANDALONE_LONG_DATE_RE = re.compile(
    r"^\s*((" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2}))\s*$"
)

# Review period example: "February 1, 2023 to January 31, 2024"
REVIEW_PERIOD_RE = re.compile(
    r"("
    r"(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2})"
    r"\s+to\s+"
    r"(" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2})"
    r")"
)

# "as of March 20, 2024" (we replace only the date portion)
AS_OF_DATE_RE = re.compile(
    r"(as of\s+)((" + "|".join(MONTHS) + r")\s+(\d{1,2}),\s+(20\d{2}))",
    re.IGNORECASE
)


# -----------------------------
# General helpers
# -----------------------------
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def _parse_rollforward_date(iso_yyyy_mm_dd: str) -> date:
    return datetime.strptime(iso_yyyy_mm_dd, "%Y-%m-%d").date()


def _format_long_date(d: date) -> str:
    # "January 15, 2025" (remove leading zero day)
    return d.strftime("%B %d, %Y").replace(" 0", " ")


def _add_one_month_same_day(d: date) -> date:
    year = d.year
    month = d.month + 1
    if month == 13:
        month = 1
        year += 1

    last_day = calendar.monthrange(year, month)[1]
    day = min(d.day, last_day)
    return date(year, month, day)


def _format_currency(user_text: str) -> str:
    s = (user_text or "").strip()
    if not s:
        raise ValueError("New amount is empty.")

    cleaned = s.replace("$", "").replace(" ", "")
    if not re.fullmatch(r"[\d,]+(\.\d{2})?", cleaned):
        raise ValueError("Amount must look like 12500, 12,500, $12,500, or 12500.00")

    has_cents = re.search(r"\.\d{2}$", cleaned) is not None
    numeric = float(cleaned.replace(",", ""))

    if has_cents:
        return f"${numeric:,.2f}"
    return f"${numeric:,.0f}"


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def build_initial_output_basename_from_uploaded_filename(uploaded_filename_raw: str, rollforward_year: int) -> str:
    """
    Initial invoice output naming:
      "[Client Name] - 202X SOC Audit - Invoice 1"
    Uses original uploaded filename for the base, but rolls forward the year and forces Invoice 1.
    """
    base = os.path.splitext(os.path.basename(uploaded_filename_raw))[0]
    base = base.replace("_", " ")
    base = re.sub(r"\s*-\s*", " - ", base)
    base = re.sub(r"\s+", " ", base).strip()

    year_pat = re.compile(r"(?<!\d)20\d{2}(?!\d)")
    if year_pat.search(base):
        base = year_pat.sub(str(rollforward_year), base, count=1)
    else:
        base = f"{base} - {rollforward_year}"

    invoice_pat = re.compile(r"(Invoice)[\s_]+(\d+)", re.IGNORECASE)
    if invoice_pat.search(base):
        base = invoice_pat.sub("Invoice 1", base, count=1)
    else:
        base = f"{base} - Invoice 1"

    base = re.sub(r"\s+", " ", base).strip()
    return base


def build_final_output_basename_from_uploaded_filename(uploaded_filename_raw: str) -> str:
    """
    Final invoice output naming:
      Keep the ORIGINAL uploaded filename base, but change "Invoice 1" -> "Invoice 2".
      Do NOT roll-forward the year in the filename (per your rule #5).
    """
    base = os.path.splitext(os.path.basename(uploaded_filename_raw))[0]
    base = base.replace("_", " ")
    base = re.sub(r"\s*-\s*", " - ", base)
    base = re.sub(r"\s+", " ", base).strip()

    invoice_pat = re.compile(r"(Invoice)[\s_]+(\d+)", re.IGNORECASE)
    if invoice_pat.search(base):
        base = invoice_pat.sub("Invoice 2", base, count=1)
    else:
        base = f"{base} - Invoice 2"

    base = re.sub(r"\s+", " ", base).strip()
    return base


def build_mgmt_rep_output_filename(raw_uploaded_filename: str) -> str:
    """
    Mgmt Rep output naming rule:
      "Change nothing about the file name except +1 to the year included in the originally uploaded file name."
    Example:
      "Management Rep Letter - 2024 Temporary Housing Directory.docx" -> "... 2025 ... .docx"
    We treat the FIRST 20xx token as the year to increment.
    """
    original_base = os.path.splitext(os.path.basename(raw_uploaded_filename))[0]
    ext = ".docx"  # output is always docx

    year_pat = re.compile(r"(?<!\d)(20\d{2})(?!\d)")
    m = year_pat.search(original_base)
    if not m:
        # If no year is present, do not invent one — keep filename identical and just force .docx
        return f"{original_base}{ext}"

    old_year = int(m.group(1))
    new_year = old_year + 1
    new_base = year_pat.sub(str(new_year), original_base, count=1)

    return f"{new_base}{ext}"


def _sanitize_filename_for_windows(name: str) -> str:
    safe = re.sub(r'[<>:"/\\|?*\n\r\t]', " ", name).strip()
    safe = re.sub(r"\s+", " ", safe)
    return safe


# -----------------------------
# Run-preserving text editing
# -----------------------------
def _get_runs_text(paragraph) -> str:
    return "".join(r.text for r in paragraph.runs)


def _build_run_spans(paragraph) -> List[Tuple[int, int]]:
    spans = []
    idx = 0
    for r in paragraph.runs:
        start = idx
        idx += len(r.text or "")
        end = idx
        spans.append((start, end))
    return spans


def _find_run_index_at_char(paragraph, char_index: int) -> Optional[int]:
    spans = _build_run_spans(paragraph)
    for i, (s, e) in enumerate(spans):
        if s <= char_index < e:
            return i
    if spans and char_index == spans[-1][1]:
        return len(spans) - 1
    return None


def _replace_span_in_runs(paragraph, start: int, end: int, replacement: str) -> bool:
    if start >= end:
        return False

    runs = paragraph.runs
    if not runs:
        return False

    spans = _build_run_spans(paragraph)

    start_run = None
    end_run = None
    for i, (s, e) in enumerate(spans):
        if start_run is None and s <= start < e:
            start_run = i
        if s < end <= e:
            end_run = i
            break

    if start_run is None:
        return False
    if end_run is None:
        end_run = len(runs) - 1

    s_start, _ = spans[start_run]
    e_start, _ = spans[end_run]
    start_off = start - s_start
    end_off = end - e_start

    before = runs[start_run].text[:start_off]
    after = runs[end_run].text[end_off:]

    runs[start_run].text = before + replacement + after

    for j in range(start_run + 1, end_run + 1):
        runs[j].text = ""

    return True


def _replace_whole_word_in_runs(paragraph, pattern: re.Pattern, repl_func) -> bool:
    changed = False
    for r in paragraph.runs:
        if not r.text:
            continue
        new_text = pattern.sub(repl_func, r.text)
        if new_text != r.text:
            r.text = new_text
            changed = True
    return changed


def _preserve_case(match: re.Match, lower: str, title: str, upper: str) -> str:
    w = match.group(0)
    if w.isupper():
        return upper
    if w[:1].isupper():
        return title
    return lower


def _preserve_case_final_to_initial(match: re.Match) -> str:
    return _preserve_case(match, lower="initial", title="Initial", upper="INITIAL")


def _preserve_case_initial_to_final(match: re.Match) -> str:
    return _preserve_case(match, lower="final", title="Final", upper="FINAL")


def _find_label_literal(full_text: str, label_regex: re.Pattern) -> Optional[str]:
    """
    Find the label portion in full_text and return it exactly as it appears (preserves casing),
    e.g., "Due Date:".
    """
    m = label_regex.search(full_text)
    if not m:
        return None
    return m.group(0)


def _replace_value_after_label_preserve_bold(paragraph, label_literal: str, new_value: str) -> bool:
    """
    Replace everything after label_literal (e.g., "DATE:" or "Due Date:") with new_value,
    preserving bold label / non-bold value styling.
    """
    full = _get_runs_text(paragraph)
    label_pos = full.find(label_literal)
    if label_pos == -1:
        return False

    start = label_pos + len(label_literal)
    while start < len(full) and full[start] == " ":
        start += 1
    end = len(full)

    label_run_idx = _find_run_index_at_char(paragraph, label_pos)
    if label_run_idx is None:
        return False

    value_run_idx = None
    for i in range(label_run_idx + 1, len(paragraph.runs)):
        run = paragraph.runs[i]
        if run.bold is False or run.bold is None:
            value_run_idx = i
            break

    if value_run_idx is None:
        return _replace_span_in_runs(paragraph, start, end, new_value)

    spans = _build_run_spans(paragraph)
    vs, _ = spans[value_run_idx]
    replace_start = max(start, vs)
    return _replace_span_in_runs(paragraph, replace_start, end, new_value)


def _replace_by_regex_on_full_text_preserve_runs(paragraph, regex: re.Pattern, repl_func) -> bool:
    """
    If regex has capturing groups, this function replaces span(1).
    If regex has no groups, it replaces the full match span.
    """
    full = _get_runs_text(paragraph)
    matches = list(regex.finditer(full))
    if not matches:
        return False

    changed = False
    for m in reversed(matches):
        start, end = m.span(1) if m.lastindex else m.span()
        replacement = repl_func(m)
        if _replace_span_in_runs(paragraph, start, end, replacement):
            changed = True

    return changed


def _replace_by_regex_group_preserve_runs(paragraph, regex: re.Pattern, group_index: int, repl_func) -> bool:
    """
    Replace a specific capturing group span (e.g., group 2) while preserving runs.
    Useful when you want to keep surrounding text like "as of ".
    """
    full = _get_runs_text(paragraph)
    matches = list(regex.finditer(full))
    if not matches:
        return False

    changed = False
    for m in reversed(matches):
        start, end = m.span(group_index)
        replacement = repl_func(m)
        if _replace_span_in_runs(paragraph, start, end, replacement):
            changed = True

    return changed


# -----------------------------
# LibreOffice conversion
# -----------------------------
def _get_soffice_cmd() -> List[str]:
    env_path = os.environ.get("LIBREOFFICE_PATH", "").strip()
    if env_path:
        return [env_path]

    soffice = shutil.which("soffice")
    if soffice:
        return [soffice]

    win_default = r"C:\Program Files\LibreOffice\program\soffice.exe"
    if os.path.exists(win_default):
        return [win_default]

    raise RuntimeError(
        "LibreOffice not found. Install LibreOffice or set LIBREOFFICE_PATH "
        "to the full path of soffice.exe/soffice."
    )


def convert_to_docx(input_path: str, out_dir: str) -> str:
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".docx":
        return input_path
    if ext != ".doc":
        raise ValueError("convert_to_docx only accepts .doc or .docx")

    cmd = _get_soffice_cmd() + [
        "--headless", "--nologo", "--nolockcheck", "--nodefault", "--norestore",
        "--convert-to", "docx",
        "--outdir", out_dir,
        input_path
    ]

    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"DOC→DOCX conversion failed:\n{res.stderr or res.stdout}")

    base = os.path.splitext(os.path.basename(input_path))[0]
    converted = os.path.join(out_dir, base + ".docx")
    if not os.path.exists(converted) or os.path.getsize(converted) == 0:
        raise RuntimeError("DOC→DOCX conversion did not produce a valid .docx output.")
    return converted


def convert_docx_to_pdf(input_docx_path: str, out_dir: str) -> str:
    if os.path.splitext(input_docx_path)[1].lower() != ".docx":
        raise ValueError("convert_docx_to_pdf requires a .docx input")

    cmd = _get_soffice_cmd() + [
        "--headless", "--nologo", "--nolockcheck", "--nodefault", "--norestore",
        "--convert-to", "pdf",
        "--outdir", out_dir,
        input_docx_path
    ]

    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        raise RuntimeError(f"DOCX→PDF conversion failed:\n{res.stderr or res.stdout}")

    base = os.path.splitext(os.path.basename(input_docx_path))[0]
    pdf_path = os.path.join(out_dir, base + ".pdf")
    if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
        raise RuntimeError("DOCX→PDF conversion did not produce a valid .pdf output.")
    return pdf_path


# -----------------------------
# Rollforward Rules (Initial Invoice)
# -----------------------------
def apply_initial_invoice_rules(
    doc: Document,
    rollforward_dt: date,
    amount_mode: str,
    new_amount_text: Optional[str]
) -> None:
    rf_date_str = _format_long_date(rollforward_dt)
    due_dt = _add_one_month_same_day(rollforward_dt)
    due_date_str = _format_long_date(due_dt)
    target_year = rollforward_dt.year

    new_amount_formatted = None
    if amount_mode == "changed":
        new_amount_formatted = _format_currency(new_amount_text or "")

    # Update title inside the document to Invoice 1
    for p in _iter_all_paragraphs(doc):
        if TITLE_HINT_RE.search(_get_runs_text(p)):
            full = _get_runs_text(p)
            parts = full.split(" - ", 1)
            if len(parts) == 2:
                client_name = parts[0].strip()
                new_title = f"{client_name} - {target_year} SOC Audit - Invoice 1"
                if p.runs:
                    p.runs[0].text = new_title
                    for r in p.runs[1:]:
                        r.text = ""
            break

    for p in _iter_all_paragraphs(doc):
        full = _get_runs_text(p)
        if not full.strip():
            continue

        # DATE:
        if DATE_LINE_RE.match(full.strip()):
            _replace_value_after_label_preserve_bold(p, "DATE:", rf_date_str)
            continue

        # Due Date:
        if DUE_DATE_LINE_RE.match(full.strip()):
            label_literal = _find_label_literal(full, re.compile(r"due date:", re.IGNORECASE))
            if label_literal:
                _replace_value_after_label_preserve_bold(p, label_literal, due_date_str)
            continue

        # final -> initial (word only)
        _replace_whole_word_in_runs(p, FINAL_WORD_RE, _preserve_case_final_to_initial)

        # year before "SOC Audit engagement" +1
        _replace_by_regex_on_full_text_preserve_runs(
            p,
            ENGAGEMENT_YEAR_RE,
            lambda m: str(int(m.group(1)) + 1)
        )

        # Professional Fees amount (optional)
        if new_amount_formatted:
            full2 = _get_runs_text(p)
            m = PRO_FEES_RE.search(full2)
            if m:
                s_start, s_end = m.span(2)
                _replace_span_in_runs(p, s_start, s_end, new_amount_formatted)

        # Invoice number update (digits 3-4 >= 10): +2 and last two "2" + random(2/4/6/8)
        full3 = _get_runs_text(p)
        matches = list(INVOICE_NUMBER_RE.finditer(full3))
        if matches:
            for mm in reversed(matches):
                six = mm.group(2)
                two_digit = int(six[2:4])
                if two_digit < 10:
                    continue

                new_two = (two_digit + 2) % 100
                new_two_str = f"{new_two:02d}"

                candidates = ["2", "4", "6", "8"]
                pick = candidates[(uuid.uuid4().int % len(candidates))]
                last_two = "2" + pick

                new_six = six[:2] + new_two_str + last_two

                s_start, s_end = mm.span(2)
                _replace_span_in_runs(p, s_start, s_end, new_six)


# -----------------------------
# Rollforward Rules (Final Invoice)
# -----------------------------
def apply_final_invoice_rules(doc: Document, rollforward_dt: date) -> None:
    """
    Final invoice rules:
      1) Invoice number: last 2 digits become "8" + random(2/4/6/8)
      2) DATE: set to calendar selection
      3) Replace "initial" -> "final" (whole word), preserve case
      4) Due date: one month after selected date, same day
      5) Preserve formatting by editing runs (no paragraph.text assignments)
    """
    rf_date_str = _format_long_date(rollforward_dt)
    due_dt = _add_one_month_same_day(rollforward_dt)
    due_date_str = _format_long_date(due_dt)

    # (Optional) Update title inside the document to Invoice 2
    for p in _iter_all_paragraphs(doc):
        if TITLE_HINT_RE.search(_get_runs_text(p)):
            full = _get_runs_text(p)
            parts = full.split(" - ", 1)
            if len(parts) == 2:
                client_name = parts[0].strip()
                m_year = re.search(r"(?<!\d)(20\d{2})(?!\d)", full)
                year_in_title = m_year.group(1) if m_year else str(rollforward_dt.year)
                new_title = f"{client_name} - {year_in_title} SOC Audit - Invoice 2"
                if p.runs:
                    p.runs[0].text = new_title
                    for r in p.runs[1:]:
                        r.text = ""
            break

    for p in _iter_all_paragraphs(doc):
        full = _get_runs_text(p)
        if not full.strip():
            continue

        # Rule #2: DATE:
        if DATE_LINE_RE.match(full.strip()):
            _replace_value_after_label_preserve_bold(p, "DATE:", rf_date_str)
            continue

        # Rule #4: Due Date:
        if DUE_DATE_LINE_RE.match(full.strip()):
            label_literal = _find_label_literal(full, re.compile(r"due date:", re.IGNORECASE))
            if label_literal:
                _replace_value_after_label_preserve_bold(p, label_literal, due_date_str)
            continue

        # Rule #3: initial -> final
        _replace_whole_word_in_runs(p, INITIAL_WORD_RE, _preserve_case_initial_to_final)

        # Rule #1: Invoice number last 2 digits -> "8" + random(2/4/6/8)
        full3 = _get_runs_text(p)
        matches = list(INVOICE_NUMBER_RE.finditer(full3))
        if matches:
            for mm in reversed(matches):
                six = mm.group(2)

                candidates = ["2", "4", "6", "8"]
                pick = candidates[(uuid.uuid4().int % len(candidates))]
                last_two = "8" + pick

                new_six = six[:4] + last_two

                s_start, s_end = mm.span(2)
                _replace_span_in_runs(p, s_start, s_end, new_six)


# -----------------------------
# Rollforward Rules (Management Rep Letter)
# -----------------------------
def apply_mgmt_rep_letter_rules(doc: Document, selected_dt: date) -> None:
    """
    Mgmt Rep Letter rules:
      1) Replace top-of-letter date (standalone "Month Day, Year") with selected date
         AND replace date portion after "as of " with selected date.
      2) Roll forward ALL review periods "Month Day, YYYY to Month Day, YYYY" by +1 year.
      3) Preserve formatting by editing runs only (no paragraph.text assignments).
    """
    selected_str = _format_long_date(selected_dt)

    for p in _iter_all_paragraphs(doc):
        full = _get_runs_text(p)

        if not full or not full.strip():
            continue

        # Rule #1a: Standalone date line (top of letter style)
        m_standalone = STANDALONE_LONG_DATE_RE.match(full)
        if m_standalone:
            # Replace only the date span (group 1 is the full date)
            date_text = m_standalone.group(1)
            idx = full.find(date_text)
            if idx != -1:
                _replace_span_in_runs(p, idx, idx + len(date_text), selected_str)
            continue

        # Rule #1b: "as of <DATE>" (replace only the date group)
        _replace_by_regex_group_preserve_runs(
            p,
            AS_OF_DATE_RE,
            group_index=2,
            repl_func=lambda m: selected_str
        )

        # Rule #2: Review periods by +1 year (replace the whole match (group 1))
        def _roll_period(m: re.Match) -> str:
            # groups: 2 month1, 3 day1, 4 year1, 5 month2, 6 day2, 7 year2
            month1 = m.group(2)
            day1 = m.group(3)
            year1 = int(m.group(4)) + 1
            month2 = m.group(5)
            day2 = m.group(6)
            year2 = int(m.group(7)) + 1
            return f"{month1} {int(day1)}, {year1} to {month2} {int(day2)}, {year2}"

        _replace_by_regex_on_full_text_preserve_runs(p, REVIEW_PERIOD_RE, _roll_period)


# -----------------------------
# Routes
# -----------------------------
@app.route("/", methods=["GET"])
def home():
    return render_template("home.html")


@app.route("/invoice", methods=["GET"])
def invoice_page():
    return render_template("invoice.html")


@app.route("/mgmt-rep", methods=["GET"])
def mgmt_rep_page():
    return render_template("mgmt_rep.html")


@app.route("/rollforward", methods=["POST"])
def rollforward():
    invoice_type = request.form.get("invoice_type", "initial").strip().lower()
    rollforward_date_raw = (request.form.get("rollforward_date") or "").strip()

    # Only used for Initial invoice
    amount_mode = (request.form.get("amount_mode") or "same").strip().lower()
    new_amount = (request.form.get("new_amount") or "").strip()

    file = request.files.get("invoice_file")

    if not file or file.filename.strip() == "":
        return "No file uploaded.", 400
    if not allowed_file(file.filename):
        return f"File type not allowed. Allowed: {sorted(ALLOWED_EXTENSIONS)}", 400
    if not rollforward_date_raw:
        return "Rollforward Date is required.", 400

    try:
        rollforward_dt = _parse_rollforward_date(rollforward_date_raw)
    except Exception:
        return "Invalid Rollforward Date format.", 400

    # Keep raw filename for output naming (preserves spaces), but use secure for saving
    raw_uploaded_name = file.filename
    safe_uploaded_name = secure_filename(file.filename)

    upload_id = uuid.uuid4().hex[:10]
    uploaded_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{upload_id}__{safe_uploaded_name}")
    file.save(uploaded_path)

    try:
        # Convert to DOCX if needed
        processing_docx_path = convert_to_docx(uploaded_path, app.config["UPLOAD_FOLDER"])
        doc = Document(processing_docx_path)

        if invoice_type == "initial":
            if amount_mode not in {"same", "changed"}:
                return "Invalid Invoice Amount selection.", 400
            if amount_mode == "changed" and not new_amount:
                return "New Amount is required when Invoice Amount = Has changed.", 400

            apply_initial_invoice_rules(
                doc=doc,
                rollforward_dt=rollforward_dt,
                amount_mode=amount_mode,
                new_amount_text=new_amount if amount_mode == "changed" else None
            )

            # Output name rolls forward year and forces Invoice 1
            out_base = build_initial_output_basename_from_uploaded_filename(raw_uploaded_name, rollforward_dt.year)
            safe_base = _sanitize_filename_for_windows(out_base)

            # Save DOCX
            out_docx_path = os.path.join(app.config["OUTPUT_FOLDER"], f"{safe_base}.docx")
            doc.save(out_docx_path)

            # Convert DOCX -> PDF
            out_pdf_path = convert_docx_to_pdf(out_docx_path, app.config["OUTPUT_FOLDER"])

            # Zip both
            zip_name = f"{safe_base} (DOCX+PDF).zip"
            zip_path = os.path.join(app.config["OUTPUT_FOLDER"], zip_name)

            with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
                z.write(out_docx_path, arcname=os.path.basename(out_docx_path))
                z.write(out_pdf_path, arcname=os.path.basename(out_pdf_path))

            return send_file(zip_path, as_attachment=True, download_name=zip_name)

        elif invoice_type == "final":
            apply_final_invoice_rules(doc, rollforward_dt)

            # Output name keeps original, only Invoice 1 -> Invoice 2
            out_base = build_final_output_basename_from_uploaded_filename(raw_uploaded_name)
            safe_base = _sanitize_filename_for_windows(out_base)

            out_docx_path = os.path.join(app.config["OUTPUT_FOLDER"], f"{safe_base}.docx")
            doc.save(out_docx_path)

            # Rule #5: DOCX only
            return send_file(out_docx_path, as_attachment=True, download_name=os.path.basename(out_docx_path))

        else:
            return "Invalid invoice_type. Use 'initial' or 'final'.", 400

    except Exception as e:
        return f"Rollforward failed: {str(e)}", 500


@app.route("/rollforward/mgmt-rep", methods=["POST"])
def rollforward_mgmt_rep():
    rollforward_date_raw = (request.form.get("rollforward_date") or "").strip()
    file = request.files.get("rep_file")

    if not file or file.filename.strip() == "":
        return "No file uploaded.", 400
    if not allowed_file(file.filename):
        return f"File type not allowed. Allowed: {sorted(ALLOWED_EXTENSIONS)}", 400
    if not rollforward_date_raw:
        return "Rollforward Date is required.", 400

    try:
        selected_dt = _parse_rollforward_date(rollforward_date_raw)
    except Exception:
        return "Invalid Rollforward Date format.", 400

    raw_uploaded_name = file.filename
    safe_uploaded_name = secure_filename(file.filename)

    upload_id = uuid.uuid4().hex[:10]
    uploaded_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{upload_id}__{safe_uploaded_name}")
    file.save(uploaded_path)

    try:
        # Convert to DOCX if needed
        processing_docx_path = convert_to_docx(uploaded_path, app.config["UPLOAD_FOLDER"])
        doc = Document(processing_docx_path)

        # Apply mgmt rep rules
        apply_mgmt_rep_letter_rules(doc, selected_dt)

        # Output filename: same as uploaded except first year +1, docx only
        out_filename = build_mgmt_rep_output_filename(raw_uploaded_name)
        out_filename_safe = _sanitize_filename_for_windows(out_filename)

        out_docx_path = os.path.join(app.config["OUTPUT_FOLDER"], out_filename_safe)
        doc.save(out_docx_path)

        return send_file(out_docx_path, as_attachment=True, download_name=os.path.basename(out_docx_path))

    except Exception as e:
        return f"Mgmt Rep rollforward failed: {str(e)}", 500


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)